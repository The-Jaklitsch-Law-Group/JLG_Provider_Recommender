import streamlit as st
import pandas as pd
import numpy as np
from geopy.geocoders import Nominatim
from geopy.extra.rate_limiter import RateLimiter
from geopy.distance import great_circle
from geopy.exc import GeocoderUnavailable
import io
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

# --- Custom Functions ---
def recommend_provider(provider_df, alpha=0.5, beta=0.5):
    """Return the best provider and scored DataFrame, prioritizing preferred providers, then lowest blended score."""
    df = provider_df.copy()
    df = df[df['Distance (miles)'].notnull() & df['Rank'].notnull()]
    if df.empty:
        return None, None
    # Prioritize preferred providers: filter to preferred if any exist
    preferred_df = df[df['Preferred'] == 1]
    if not preferred_df.empty:
        df = preferred_df
    df['norm_rank'] = (df['Rank'] - df['Rank'].min()) / (df['Rank'].max() - df['Rank'].min())
    df['norm_dist'] = (df['Distance (miles)'] - df['Distance (miles)'].min()) / (df['Distance (miles)'].max() - df['Distance (miles)'].min())
    df['score'] = alpha * df['norm_rank'] + beta * df['norm_dist']
    best = df.sort_values(by='score').iloc[0]
    return best, df

@st.cache_data(show_spinner=True)
def geocode_providers(addresses, _geocode):
    """Geocode a list of addresses, returning lists of latitudes and longitudes."""
    lats, lons = [], []
    for addr in addresses:
        try:
            location = _geocode(addr, timeout=10)
            if location:
                lats.append(location.latitude)
                lons.append(location.longitude)
            else:
                lats.append(None)
                lons.append(None)
        except GeocoderUnavailable as e:
            lats.append(None)
            lons.append(None)
            print(f"GeocoderUnavailable for address '{addr}': {e}")
        except Exception as e:
            lats.append(None)
            lons.append(None)
            print(f"Geocoding error for address '{addr}': {e}")
    return lats, lons

def calculate_distances(user_lat, user_lon, provider_df):
    """Calculate distances in miles from user to each provider."""
    user_loc = (user_lat, user_lon)
    distances = []
    for lat, lon in zip(provider_df['Latitude'], provider_df['Longitude']):
        if pd.notnull(lat) and pd.notnull(lon):
            dist = great_circle(user_loc, (lat, lon)).miles
        else:
            dist = None
        distances.append(dist)
    return distances

# --- Data Loading ---
def load_provider_data():
    df = pd.read_excel('data/Ranked_Contacts.xlsx')
    df.columns = [col.strip() for col in df.columns]
    df['Address 1 Zip'] = df['Address 1 Zip'].apply(lambda x: str(int(x)) if pd.notnull(x) else '')
    df['Full Address'] = (
        df['Address 1 Line 1'].fillna('') + ', '
        + df['Address 1 City'].fillna('') + ', '
        + df['Address 1 State'].fillna('') + ' '
        + df['Address 1 Zip'].fillna('')
    )
    df['Full Address'] = df['Full Address'].str.replace(r',\s*,', ',', regex=True).str.replace(r',\s*$', '', regex=True)
    return df

provider_df = load_provider_data()

# --- Add temporary placeholders if needed ---
if 'Specialty' not in provider_df.columns:
    placeholder_specialties = [
        'Primary Care', 'Urgent Care', 'Chiropractor', 'Physical Therapy'
    ]
    provider_df['Specialty'] = np.random.choice(placeholder_specialties, size=len(provider_df))
if 'Phone 1' not in provider_df.columns:
    provider_df['Phone 1'] = '555-555-1234'
if 'Email 1' not in provider_df.columns:
    provider_df['Email 1'] = 'provider@example.com'
if 'Preferred' not in provider_df.columns:
    provider_df['Preferred'] = np.random.choice([1, 0], size=len(provider_df), p=[0.3, 0.7])  # 30% preferred

# --- Load API secrets for future API calls ---
LD_API_KEY = st.secrets.get('lead_docket_api_key', None)
LD_API_ENDPOINT = st.secrets.get('lead_docket_base_api_endpoint', None)

def call_external_api(payload):
    """Placeholder for making an API call using the stored API key and endpoint."""
    if not LD_API_KEY or not LD_API_ENDPOINT:
        st.warning('API key or endpoint not set in secrets. Please update .streamlit/secrets.toml.')
        return None
    # Example usage (requests must be imported if used):
    # import requests
    # headers = {"Authorization": f"Bearer {API_KEY}"}
    # response = requests.post(API_ENDPOINT, json=payload, headers=headers)
    # return response.json()
    return None  # Placeholder

st.set_page_config(page_title="Provider Recommender", page_icon=":hospital:")

st.title('Provider Recommender for Personal Injury Clients')

tabs = st.tabs(["Find Provider", "How Selection Works"])

# --- Instructions in Sidebar ---
st.sidebar.markdown("""
# Instructions:
1. Click the **New Search** button to start a new recommendation.<br>
2. Enter the client's address and preferences.<br>
3. Select a provider specialty if needed.<br>
4. Adjust the sliders to set the importance (weight) of provider rank vs. distance.<br>
5. Click <b>Find Best Provider</b> to get a recommendation. The app prioritizes the law firm's preferred providers, then considers proximity and ranking.<br>
6. The final result is contact information to direct the client to the best provider.
""", unsafe_allow_html=True)

with tabs[0]:
    # --- Layout: Form on left, results on right ---
    left_col, right_col = st.columns([.5, .5])
    with left_col:
        # The form is always present
        # Initialize variables to avoid unbound errors
        street = st.session_state.get('street', '')
        city = st.session_state.get('city', '')
        state = st.session_state.get('state', '')
        zipcode = st.session_state.get('zipcode', '')
        specialty = st.session_state.get('specialty', 'Any')
        alpha = st.session_state.get('alpha', 0.5)
        beta = st.session_state.get('beta', 0.5)
        submit = False
        specialties = sorted(provider_df['Specialty'].dropna().unique()) if 'Specialty' in provider_df.columns else []
        with st.form(key='address_form'):
            street = st.text_input('Street Address (e.g., 123 Main St)', value=street)
            city = st.text_input('City', value=city)
            state = st.text_input('State (e.g., NY)', value=state)
            zipcode = st.text_input('Zip Code', value=zipcode)
            specialty = st.selectbox('Provider Specialty (optional)', ['Any'] + specialties) if specialties else 'Any'
            col1, col2 = st.columns(2)
            with col1:
                alpha = st.slider('How important is provider quality?', min_value=0.0, max_value=1.0, value=alpha, step=0.05, help="Move right to prioritize provider quality (ranking)")
            with col2:
                beta = st.slider('How important is provider proximity?', min_value=0.0, max_value=1.0, value=beta, step=0.05, help="Move right to prioritize providers closer to the address")
            if alpha + beta != 1.0:
                total = alpha + beta
                alpha = alpha / total
                beta = beta / total
            submit = st.form_submit_button('Find Best Provider')

        st.session_state['street'] = street
        st.session_state['city'] = city
        st.session_state['state'] = state
        st.session_state['zipcode'] = zipcode
        st.session_state['alpha'] = alpha
        st.session_state['beta'] = beta
        st.session_state['specialty'] = specialty

    with right_col:
        # --- Geocoding Setup ---
        geolocator = Nominatim(user_agent="provider_recommender")
        geocode = RateLimiter(geolocator.geocode, min_delay_seconds=2, max_retries=3)
        with st.spinner('Geocoding providers and calculating recommendations...'):
            provider_lats, provider_lons = geocode_providers(provider_df['Full Address'], geocode)
            provider_df['Latitude'] = provider_lats
            provider_df['Longitude'] = provider_lons

            # --- Content for Results ---
            if submit:
                user_full_address = f"{street}, {city}, {state} {zipcode}".strip(', ')
                user_lat, user_lon = None, None
                try:
                    user_location = geocode(user_full_address, timeout=10)
                    if not user_location and street:
                        street_simple = street.split(',')[0].split(' Apt')[0].split(' Suite')[0]
                        user_location = geocode(street_simple, timeout=10)
                    if not user_location:
                        if city and state:
                            user_location = geocode(f"{city}, {state}", timeout=10)
                        elif zipcode:
                            user_location = geocode(zipcode, timeout=10)
                    if user_location:
                        user_lat, user_lon = user_location.latitude, user_location.longitude
                        st.session_state['user_lat'] = user_lat
                        st.session_state['user_lon'] = user_lon
                    else:
                        st.error("Could not geocode your address. Please check and try again.")
                except GeocoderUnavailable as e:
                    st.error(f"Geocoding service unavailable: {e}")
                except Exception as e:
                    st.error(f"Geocoding error: {e}")

                if user_lat is not None and user_lon is not None:
                    filtered_df = provider_df.copy()
                    if specialty != 'Any' and 'Specialty' in filtered_df.columns:
                        filtered_df = filtered_df[filtered_df['Specialty'] == specialty]
                    filtered_df['Distance (miles)'] = calculate_distances(user_lat, user_lon, filtered_df)

                    best, scored_df = recommend_provider(filtered_df, alpha=alpha, beta=beta)
                    st.subheader('Best Provider Recommendation')
                    if best is not None and isinstance(scored_df, pd.DataFrame):
                        st.markdown(f"👨‍⚕️ **Name:** {best['Full Name']}")
                        address_for_url = best['Full Address'].replace(' ', '+')
                        maps_url = f"https://www.google.com/maps/search/?api=1&query={address_for_url}"
                        st.markdown(f"🏥 **Address:** [{best['Full Address']}]({maps_url})")
                        st.markdown(f"📞 **Phone:** {best['Phone 1']}")
                        st.markdown(f"📧 **Email:** {best['Email 1']}")
                        st.markdown(f"🏥 **Specialty:** {best['Specialty']}")
                        if best.get('Preferred', 0) == 1:
                            st.markdown(f"<span style='color: green; font-weight: bold;'>✅ Preferred Provider</span>", unsafe_allow_html=True)
                        st.write('Top 5 providers by blended score:')
                        st.dataframe(scored_df[['Full Name', 'Full Address', 'Distance (miles)', 'Rank', 'score', 'Preferred']].sort_values(by='score').head())

                        # --- Export Buttons ---
                        def get_word_bytes(best):
                            doc = Document()
                            doc.add_heading('Recommended Provider', 0)
                            doc.add_paragraph(f"Name: {best['Full Name']}")
                            doc.add_paragraph(f"Address: {best['Full Address']}")
                            doc.add_paragraph(f"Phone: {best['Phone 1']}")
                            doc.add_paragraph(f"Email: {best['Email 1']}")
                            doc.add_paragraph(f"Specialty: {best['Specialty']}")
                            if best.get('Preferred', 0) == 1:
                                doc.add_paragraph("Preferred Provider")
                            buffer = io.BytesIO()
                            doc.save(buffer)
                            buffer.seek(0)
                            return buffer

                        def get_pdf_bytes(best):
                            buffer = io.BytesIO()
                            c = canvas.Canvas(buffer, pagesize=letter)
                            c.setFont("Helvetica-Bold", 16)
                            c.drawString(72, 720, "Recommended Provider")
                            c.setFont("Helvetica", 12)
                            y = 700
                            c.drawString(72, y, f"Name: {best['Full Name']}")
                            y -= 20
                            c.drawString(72, y, f"Address: {best['Full Address']}")
                            y -= 20
                            c.drawString(72, y, f"Phone: {best['Phone 1']}")
                            y -= 20
                            c.drawString(72, y, f"Email: {best['Email 1']}")
                            y -= 20
                            c.drawString(72, y, f"Specialty: {best['Specialty']}")
                            y -= 20
                            if best.get('Preferred', 0) == 1:
                                c.drawString(72, y, "Preferred Provider")
                            c.save()
                            buffer.seek(0)
                            return buffer

                        word_bytes = get_word_bytes(best)
                        st.download_button(
                            label="Export as Word",
                            data=word_bytes,
                            file_name="recommended_provider.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                        pdf_bytes = get_pdf_bytes(best)
                        st.download_button(
                            label="Export as PDF",
                            data=pdf_bytes,
                            file_name="recommended_provider.pdf",
                            mime="application/pdf"
                        )

                        # --- Rationale for Selection ---
                        with st.expander('Why was this provider selected?', expanded=False):
                            rationale = []
                            if best.get('Preferred', 0) == 1:
                                rationale.append('This provider is a **Preferred Provider** for the law firm, which means they are trusted and have a strong track record with our clients.')
                            else:
                                rationale.append('This provider is not marked as preferred, but was selected based on a balance of proximity and ranking.')
                            rationale.append(f"")
                            rationale.append(f"* **Distance** from the address is **{best['Distance (miles)']:.2f} miles**.")
                            rationale.append(f"")
                            rationale.append(f"* Selected specialty is **{best['Specialty']}**.")
                            rationale.append(f"")
                            rationale.append(f"* Provider's rank is **{best['Rank']}** (lower is better).")
                            rationale.append(f"")
                            rationale.append(f"The final score is a blend of normalized rank and distance, using your chosen weights: **Rank weight = {alpha:.2f}**, **Distance weight = {beta:.2f}**.")
                            rationale.append(f"The provider with the lowest blended score was recommended.")
                            st.markdown('<br>'.join(rationale), unsafe_allow_html=True)
                    else:
                        st.warning('No providers with both rank and distance available.')


with tabs[1]:
    # --- High-Level Overview of Selection Process ---
    st.markdown("""
    ### How Provider Selection Works
    - This app is designed for personal injury law firms to help new clients identify healthcare providers.
    - The law firm's **preferred providers** are prioritized in recommendations.
    - **Step 1:** The client's address is geocoded (converted to latitude/longitude). If the full address fails, the app tries less specific versions (street only, city/state, or zip).
    - **Step 2:** The distance from the client to each provider is calculated.
    - **Step 3:** Each provider has a pre-assigned rank (lower is better).
    - **Step 4:** The app calculates a **blended score** for each provider:
        - The score is a weighted sum: `score = weight_rank * normalized_rank + weight_distance * normalized_distance`
        - Both rank and distance are normalized (scaled between 0 and 1) so they are comparable.
        - The weights are set by you using the sliders above.
        - Lower scores mean a better balance of proximity and provider quality.
    - **Step 5:** If any preferred providers are available, only those are considered for the final recommendation.
    - **Step 6:** The provider with the lowest blended score is recommended as the best option.
    - The final product is clear contact information to direct the client to the best provider.
    - You will also see the top 5 providers by blended score for comparison.
    """)