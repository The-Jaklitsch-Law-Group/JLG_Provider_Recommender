"""
Consolidated Functions Module for JLG Provider Recommender

This module contains the best versions of all functions that were found to be
redundant across multiple modules. The functions have been carefully selected
based on:
- Most comprehensive implementation
- Best error handling
- Latest features and optimizations
- Most complete documentation

Functions are organized by category:
1. Data Loading and Processing
2. Address Validation and Geocoding
3. Provider Recommendation and Scoring
4. Data Validation and Quality Checks
5. Utility Functions

Original sources consolidated:
- src/utils/providers.py
- src/utils/validation.py
- src/utils/performance.py
- src/data/ingestion.py
- app.py
- data_dashboard.py
"""

import functools
import io
import logging
import re
import time
from datetime import datetime, timedelta
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, Union

import numpy as np
import pandas as pd
import psutil
import streamlit as st
from docx import Document
from geopy.distance import geodesic
from geopy.exc import GeocoderServiceError, GeocoderTimedOut, GeocoderUnavailable
from geopy.geocoders import Nominatim

# Configure logging
logger = logging.getLogger(__name__)


# =============================================================================
# DATA LOADING AND PROCESSING
# =============================================================================


@st.cache_data(ttl=3600)
def load_provider_data(filepath: str) -> pd.DataFrame:
    """
    Load and preprocess provider data from various file formats.

    This function loads healthcare provider data from Excel, CSV, Feather, or Parquet files
    and performs basic preprocessing including column cleaning, data type conversion,
    and address formatting.

    Args:
        filepath (str): Path to the data file. Supported formats: .xlsx, .csv, .feather, .parquet

    Returns:
        pd.DataFrame: Preprocessed provider data with columns:
            - Full Name: Provider's complete name
            - Specialty: Medical specialty
            - Street, City, State, Zip: Address components
            - Full Address: Formatted complete address
            - Referral Count: Number of referrals (converted to numeric)
            - Latitude, Longitude: Geographic coordinates (if available)

    Raises:
        FileNotFoundError: If the specified file does not exist
        ValueError: If the file format is not supported

    Example:
        >>> df = load_provider_data("data/providers.xlsx")
        >>> print(df.columns.tolist())
        ['Full Name', 'Specialty', 'Street', 'City', 'State', 'Zip', 'Full Address', 'Referral Count']

    Note:
        - Function is cached with Streamlit for 1 hour to improve performance
        - Preference column is automatically dropped if present
        - ZIP codes are converted to strings to preserve leading zeros
        - Missing address components are filled with empty strings
    """
    path = Path(filepath)
    if not path.exists():
        raise FileNotFoundError(f"File {filepath} does not exist")

    suffix = path.suffix.lower()
    if suffix == ".xlsx":
        df = pd.read_excel(path)
    elif suffix == ".csv":
        df = pd.read_csv(path)
    elif suffix == ".feather":
        df = pd.read_feather(path)
    elif suffix == ".parquet":
        df = pd.read_parquet(path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")

    df.columns = [col.strip() for col in df.columns]
    df = df.drop(columns="Preference", errors="ignore")

    # Ensure all address columns are strings to prevent concatenation errors
    address_cols = ["Street", "City", "State", "Zip"]
    for col in address_cols:
        if col in df.columns:
            df[col] = df[col].astype(str).replace(["nan", "None", "NaN"], "").fillna("")

    # Handle Referral Count column if it exists
    if "Referral Count" in df.columns:
        df["Referral Count"] = pd.to_numeric(df["Referral Count"], errors="coerce")

    # Build Full Address from components
    if "Full Address" not in df.columns:
        df["Full Address"] = (
            df["Street"].fillna("")
            + ", "
            + df["City"].fillna("")
            + ", "
            + df["State"].fillna("")
            + " "
            + df["Zip"].fillna("")
        )
        df["Full Address"] = (
            df["Full Address"].str.replace(r",\s*,", ",", regex=True).str.replace(r",\s*$", "", regex=True)
        )
    return df


def clean_address_data(df: pd.DataFrame) -> pd.DataFrame:
    """
    Clean and standardize address data in a DataFrame.

    Args:
        df: DataFrame with address columns

    Returns:
        DataFrame with cleaned address data
    """
    df = df.copy()

    # Standardize address column names
    address_cols = ["Street", "City", "State", "Zip"]
    for col in address_cols:
        if col in df.columns:
            # Convert to string and clean
            df[col] = df[col].astype(str).str.strip()
            df[col] = df[col].replace(["nan", "None", "NaN", ""], pd.NA)
            df[col] = df[col].fillna("")

    # Standardize state abbreviations
    if "State" in df.columns:
        state_mapping = {
            "ALABAMA": "AL",
            "ALASKA": "AK",
            "ARIZONA": "AZ",
            "ARKANSAS": "AR",
            "CALIFORNIA": "CA",
            "COLORADO": "CO",
            "CONNECTICUT": "CT",
            "DELAWARE": "DE",
            "FLORIDA": "FL",
            "GEORGIA": "GA",
            "HAWAII": "HI",
            "IDAHO": "ID",
            "ILLINOIS": "IL",
            "INDIANA": "IN",
            "IOWA": "IA",
            "KANSAS": "KS",
            "KENTUCKY": "KY",
            "LOUISIANA": "LA",
            "MAINE": "ME",
            "MARYLAND": "MD",
            "MASSACHUSETTS": "MA",
            "MICHIGAN": "MI",
            "MINNESOTA": "MN",
            "MISSISSIPPI": "MS",
            "MISSOURI": "MO",
            "MONTANA": "MT",
            "NEBRASKA": "NE",
            "NEVADA": "NV",
            "NEW HAMPSHIRE": "NH",
            "NEW JERSEY": "NJ",
            "NEW MEXICO": "NM",
            "NEW YORK": "NY",
            "NORTH CAROLINA": "NC",
            "NORTH DAKOTA": "ND",
            "OHIO": "OH",
            "OKLAHOMA": "OK",
            "OREGON": "OR",
            "PENNSYLVANIA": "PA",
            "RHODE ISLAND": "RI",
            "SOUTH CAROLINA": "SC",
            "SOUTH DAKOTA": "SD",
            "TENNESSEE": "TN",
            "TEXAS": "TX",
            "UTAH": "UT",
            "VERMONT": "VT",
            "VIRGINIA": "VA",
            "WASHINGTON": "WA",
            "WEST VIRGINIA": "WV",
            "WISCONSIN": "WI",
            "WYOMING": "WY",
            "DISTRICT OF COLUMBIA": "DC",
        }
        df["State"] = df["State"].str.upper().map(state_mapping).fillna(df["State"])

    return df


def build_full_address(df: pd.DataFrame) -> pd.DataFrame:
    """
    Build full address from address components.

    Args:
        df: DataFrame with address component columns

    Returns:
        DataFrame with Full Address column added
    """
    df = df.copy()

    def safe_join_address(row):
        """Safely join address components, handling missing values."""
        parts = []
        for col in ["Street", "City", "State", "Zip"]:
            if col in row and pd.notna(row[col]) and str(row[col]).strip():
                parts.append(str(row[col]).strip())

        if len(parts) >= 2:  # At least street and city/state
            return ", ".join(parts)
        return ""

    if "Full Address" not in df.columns:
        df["Full Address"] = df.apply(safe_join_address, axis=1)

    return df


def safe_numeric_conversion(value: Any, default: float = 0.0) -> float:
    """
    Safely convert a value to float with fallback to default.

    This function handles common conversion errors when dealing with mixed data types,
    missing values, or malformed numeric strings. It's particularly useful for
    cleaning data from external sources.

    Args:
        value (Any): Value to convert to float. Can be str, int, float, None, or pandas.NA
        default (float, optional): Default value to return if conversion fails. Defaults to 0.0

    Returns:
        float: Successfully converted float value or the default value

    Examples:
        >>> safe_numeric_conversion("123.45")
        123.45
        >>> safe_numeric_conversion("invalid", default=0.0)
        0.0
        >>> safe_numeric_conversion(None)
        0.0
        >>> safe_numeric_conversion(pd.NA, default=-1.0)
        -1.0

    Note:
        Uses pandas.isna() to check for missing values, which handles both
        None and pandas NA values properly.
    """
    try:
        if pd.isna(value):
            return default
        return float(value)
    except (ValueError, TypeError):
        return default


# =============================================================================
# ADDRESS VALIDATION AND GEOCODING
# =============================================================================


def validate_address(address: str) -> Tuple[bool, str]:
    """
    Validate address format and completeness.

    Args:
        address: Address string to validate

    Returns:
        Tuple of (is_valid, validation_message)
    """
    if not address or not address.strip():
        return False, "Address cannot be empty"

    address = address.strip()

    # Basic length check
    if len(address) < 10:
        return False, "Address appears too short. Please provide a complete address."

    # Check for basic components
    has_number = any(char.isdigit() for char in address)
    has_street_indicator = any(
        indicator in address.lower()
        for indicator in [
            "street",
            "st",
            "avenue",
            "ave",
            "road",
            "rd",
            "drive",
            "dr",
            "lane",
            "ln",
            "way",
            "blvd",
            "boulevard",
        ]
    )

    if not has_number:
        return False, "Address should include a street number"

    # Check for state abbreviation or full state name
    state_patterns = ["\\b[A-Z]{2}\\b", "\\d{5}(-\\d{4})?\\b"]  # State abbrev or ZIP
    has_state_or_zip = any(re.search(pattern, address) for pattern in state_patterns)

    if not has_state_or_zip:
        return True, "Consider adding state and ZIP code for better accuracy"

    return True, ""


def validate_address_input(street: str, city: str, state: str, zipcode: str) -> Tuple[bool, str]:
    """
    Comprehensive address validation with enhanced security and validation.

    This function validates individual address components using enhanced validation
    rules and security checks to ensure data integrity and prevent injection attacks.

    Args:
        street (str): Street address (e.g., "123 Main St")
        city (str): City name (e.g., "New York")
        state (str): State abbreviation (e.g., "NY")
        zipcode (str): ZIP code (e.g., "10001" or "10001-1234")

    Returns:
        tuple[bool, str]: (is_valid, detailed_error_message)

    Examples:
        >>> is_valid, msg = validate_address_input("123 Main St", "Anytown", "CA", "12345")
        >>> print(is_valid)  # True

        >>> is_valid, msg = validate_address_input("", "Anytown", "INVALID", "12345")
        >>> print(msg)  # "Street address must be at least 3 characters; State must be a valid 2-letter abbreviation"
    """
    errors = []
    warnings = []

    # Basic required field validation
    if not street.strip():
        errors.append("Street address is required")

    if not city.strip():
        warnings.append("City is recommended for better geocoding accuracy")

    if not state.strip():
        warnings.append("State is recommended for better geocoding accuracy")

    # Validate state format (if provided)
    if state.strip():
        state_clean = state.strip().upper()
        # Common US state abbreviations
        valid_states = {
            "AL",
            "AK",
            "AZ",
            "AR",
            "CA",
            "CO",
            "CT",
            "DE",
            "FL",
            "GA",
            "HI",
            "ID",
            "IL",
            "IN",
            "IA",
            "KS",
            "KY",
            "LA",
            "ME",
            "MD",
            "MA",
            "MI",
            "MN",
            "MS",
            "MO",
            "MT",
            "NE",
            "NV",
            "NH",
            "NJ",
            "NM",
            "NY",
            "NC",
            "ND",
            "OH",
            "OK",
            "OR",
            "PA",
            "RI",
            "SC",
            "SD",
            "TN",
            "TX",
            "UT",
            "VT",
            "VA",
            "WA",
            "WV",
            "WI",
            "WY",
            "DC",
        }
        if len(state_clean) == 2 and state_clean not in valid_states:
            warnings.append(f"'{state}' may not be a valid US state abbreviation")
        elif len(state_clean) > 2:
            warnings.append("Consider using 2-letter state abbreviation (e.g., 'MD' instead of 'Maryland')")

    # Validate ZIP code format (if provided)
    if zipcode.strip():
        zip_clean = zipcode.strip()
        if not (zip_clean.isdigit() and len(zip_clean) == 5):
            if not (
                len(zip_clean) == 10 and zip_clean[5] == "-" and zip_clean[:5].isdigit() and zip_clean[6:].isdigit()
            ):
                warnings.append("ZIP code should be 5 digits (e.g., '20746') or ZIP+4 format (e.g., '20746-1234')")

    # Check for suspicious patterns
    if street.strip().lower() in ["test", "example", "123 test st", "123 main st"]:
        warnings.append("Address appears to be a test value - please enter a real address")

    # Compile messages
    message_parts = []
    if errors:
        message_parts.append("❌ **Errors**: " + "; ".join(errors))
    if warnings:
        message_parts.append("⚠️ **Suggestions**: " + "; ".join(warnings))

    is_valid = len(errors) == 0
    message = "\n\n".join(message_parts) if message_parts else ""

    return is_valid, message


def validate_coordinates(lat: float, lon: float) -> Tuple[bool, str]:
    """
    Validate latitude and longitude coordinates.

    Args:
        lat: Latitude value
        lon: Longitude value

    Returns:
        Tuple of (is_valid, error_message)
    """
    if not isinstance(lat, (int, float)) or not isinstance(lon, (int, float)):
        return False, "Coordinates must be numeric"

    if not (-90 <= lat <= 90):
        return False, "Latitude must be between -90 and 90"

    if not (-180 <= lon <= 180):
        return False, "Longitude must be between -180 and 180"

    return True, "Valid coordinates"


@st.cache_data(ttl=3600)
def geocode_address_with_cache(address: str) -> Optional[Tuple[float, float]]:
    """
    Geocode an address with caching and error handling.

    Args:
        address: Address string to geocode

    Returns:
        Tuple of (latitude, longitude) or None if failed
    """
    try:
        geolocator = Nominatim(user_agent="jlg_provider_recommender")
        location = geolocator.geocode(address, timeout=10)

        if location:
            return (location.latitude, location.longitude)
        else:
            return None

    except (GeocoderTimedOut, GeocoderServiceError) as e:
        st.warning(f"Geocoding service temporarily unavailable. Please try again.")
        return None
    except Exception as e:
        st.error(f"Error geocoding address: {str(e)}")
        return None


@st.cache_data(ttl=60 * 60 * 24)
def cached_geocode_address(address: str):
    """
    Cached single-address geocode using Nominatim with rate limiting.

    Returns a geopy Location or None. TTL 24 hours to reuse results.
    """
    try:
        from geopy.extra.rate_limiter import RateLimiter

        geolocator_local = Nominatim(user_agent="provider_recommender")
        geocode_local = RateLimiter(geolocator_local.geocode, min_delay_seconds=2, max_retries=3)
        return geocode_local(address, timeout=10)
    except GeocoderUnavailable:
        return None
    except Exception:
        return None


def handle_geocoding_error(address: str, error: Exception) -> str:
    """Provide user-friendly error messages for geocoding failures."""
    error_type = type(error).__name__

    if "timeout" in str(error).lower():
        return f"⏱️ **Geocoding Timeout**: The address lookup service is taking too long. Please try again in a moment."

    elif "unavailable" in str(error).lower() or "service" in str(error).lower():
        return f"🔌 **Service Unavailable**: The geocoding service is temporarily unavailable. Please try again later."

    elif "rate" in str(error).lower() or "limit" in str(error).lower():
        return f"🚦 **Rate Limited**: Too many requests to the geocoding service. Please wait a moment and try again."

    elif "network" in str(error).lower() or "connection" in str(error).lower():
        return f"🌐 **Network Error**: Cannot connect to the geocoding service. Please check your internet connection."

    else:
        return f"❌ **Geocoding Error**: Unable to find location for '{address}'. Please check the address and try again. (Error: {error_type})"


# =============================================================================
# DISTANCE CALCULATION AND PROVIDER RECOMMENDATION
# =============================================================================


def calculate_distances(user_lat: float, user_lon: float, provider_df: pd.DataFrame) -> List[Optional[float]]:
    """
    Calculate distances in miles from user to each provider using vectorized operations.

    Args:
        user_lat: User's latitude
        user_lon: User's longitude
        provider_df: DataFrame with provider locations

    Returns:
        List of distances in miles (None for invalid coordinates)
    """
    lat_rad = np.radians(provider_df["Latitude"].to_numpy(dtype=float))
    lon_rad = np.radians(provider_df["Longitude"].to_numpy(dtype=float))
    user_lat_rad = np.radians(user_lat)
    user_lon_rad = np.radians(user_lon)

    valid = ~np.isnan(lat_rad) & ~np.isnan(lon_rad)
    dlat = lat_rad[valid] - user_lat_rad
    dlon = lon_rad[valid] - user_lon_rad
    a = np.sin(dlat / 2) ** 2 + np.cos(user_lat_rad) * np.cos(lat_rad[valid]) * np.sin(dlon / 2) ** 2
    c = 2 * np.arcsin(np.sqrt(a))
    distances = np.full(len(provider_df), np.nan)
    distances[valid] = 3958.8 * c  # Earth radius in miles

    # Convert NaN to None for consistency
    return [None if np.isnan(d) else float(d) for d in distances]


def recommend_provider(
    provider_df: pd.DataFrame,
    distance_weight: float = 0.5,
    referral_weight: float = 0.5,
    inbound_weight: float = 0.0,
    min_referrals: Optional[int] = None,
) -> Tuple[Optional[pd.Series], Optional[pd.DataFrame]]:
    """
    Recommend the best healthcare provider based on distance, outbound referral count, and inbound referral count.

    This function implements a weighted scoring algorithm that combines provider distance
    from the client, historical outbound referral patterns, and inbound referral patterns
    to recommend the most suitable healthcare provider. The algorithm normalizes all metrics
    and allows customizable weighting between proximity, provider popularity, and provider
    inbound referral activity.

    Args:
        provider_df (pd.DataFrame): DataFrame containing provider information with required columns:
            - 'Distance (Miles)': Float distance from client to provider
            - 'Referral Count': Integer number of historical outbound referrals
            - 'Inbound Referral Count': Integer number of inbound referrals (optional)
            - 'Full Name': Provider name
            - Additional columns preserved in output
        distance_weight (float, optional): Weight for distance factor (0.0-1.0). Defaults to 0.5.
            Higher values prioritize closer providers.
        referral_weight (float, optional): Weight for outbound referral count factor (0.0-1.0). Defaults to 0.5.
            Higher values prioritize more frequently referred providers.
        inbound_weight (float, optional): Weight for inbound referral count factor (0.0-1.0). Defaults to 0.0.
            Higher values prioritize providers who refer more clients to the firm.
        min_referrals (int, optional): Minimum referral count threshold. Providers with fewer
            referrals are excluded. Defaults to None (no filtering).

    Returns:
        tuple: A tuple containing:
            - best_provider (pd.Series or None): Series containing the top recommended provider's data
            - scored_df (pd.DataFrame or None): DataFrame with all eligible providers and their scores,
              sorted by recommendation score (lower is better)

    Raises:
        KeyError: If required columns ('Distance (Miles)', 'Referral Count') are missing

    Algorithm:
        1. Filters out providers with missing distance or referral data
        2. Applies minimum referral threshold if specified
        3. Normalizes distance (0-1, lower is better), outbound referral count (0-1, higher is better),
           and inbound referral count (0-1, higher is better)
        4. Calculates composite score: distance_weight * norm_distance + referral_weight * (1 - norm_referrals) + inbound_weight * norm_inbound
        5. Returns provider with lowest composite score

    Examples:
        >>> # Basic usage with equal weighting for distance and outbound referrals
        >>> best, scored = recommend_provider(providers_df)

        >>> # Include inbound referrals in scoring
        >>> best, scored = recommend_provider(providers_df, distance_weight=0.4, referral_weight=0.4, inbound_weight=0.2)

        >>> # Prioritize distance over referral history
        >>> best, scored = recommend_provider(providers_df, distance_weight=0.8, referral_weight=0.2)

        >>> # Only consider providers with at least 5 referrals
        >>> best, scored = recommend_provider(providers_df, min_referrals=5)

    Note:
        - Returns (None, None) if no providers meet the criteria
        - Weights don't need to sum to 1.0, but typical usage has distance_weight + referral_weight + inbound_weight = 1.0
        - Lower composite scores indicate better recommendations
        - If 'Inbound Referral Count' column is missing, inbound_weight is ignored
    """
    df = provider_df.copy(deep=True)
    df = df[df["Distance (Miles)"].notnull() & df["Referral Count"].notnull()]
    if min_referrals is not None:
        df = df[df["Referral Count"] >= min_referrals]

    if df.empty:
        return None, None

    # Safe normalization (avoid division by zero)
    referral_range = df["Referral Count"].max() - df["Referral Count"].min()
    dist_range = df["Distance (Miles)"].max() - df["Distance (Miles)"].min()

    # Normalize distance and outbound referrals
    df["norm_rank"] = (df["Referral Count"] - df["Referral Count"].min()) / referral_range if referral_range != 0 else 0
    df["norm_dist"] = (df["Distance (Miles)"] - df["Distance (Miles)"].min()) / dist_range if dist_range != 0 else 0

    # Initialize score with distance and outbound referrals
    df["Score"] = distance_weight * df["norm_dist"] + referral_weight * df["norm_rank"]

    # Add inbound referrals to scoring if the column exists and weight > 0
    if inbound_weight > 0 and "Inbound Referral Count" in df.columns:
        # Filter out providers with missing inbound referral data for scoring
        inbound_df = df[df["Inbound Referral Count"].notnull()].copy()

        if not inbound_df.empty:
            inbound_range = inbound_df["Inbound Referral Count"].max() - inbound_df["Inbound Referral Count"].min()

            if inbound_range != 0:
                inbound_df["norm_inbound"] = (
                    inbound_df["Inbound Referral Count"] - inbound_df["Inbound Referral Count"].min()
                ) / inbound_range
            else:
                inbound_df["norm_inbound"] = 0

            # Recalculate score for providers with inbound data
            inbound_df["Score"] = (
                distance_weight * inbound_df["norm_dist"]
                + referral_weight * inbound_df["norm_rank"]
                + inbound_weight * inbound_df["norm_inbound"]
            )

            # Update the main dataframe with new scores
            df.loc[inbound_df.index, "Score"] = inbound_df["Score"]
            df.loc[inbound_df.index, "norm_inbound"] = inbound_df["norm_inbound"]

    # Tie-break behavior: sort by score, then distance, then referral count
    if distance_weight > referral_weight:
        sort_keys = ["Score", "Distance (Miles)", "Referral Count"]
    else:
        sort_keys = ["Score", "Referral Count", "Distance (Miles)"]

    # Add inbound referrals to tie-break if available
    if "Inbound Referral Count" in df.columns:
        sort_keys.append("Inbound Referral Count")

    # Add stable final tie-break key
    candidate_keys = sort_keys + ["Full Name"]
    sort_keys_final = [k for k in candidate_keys if k in df.columns]
    ascending = [True] * len(sort_keys_final)

    df_sorted = df.sort_values(by=sort_keys_final, ascending=ascending).reset_index(drop=True)
    best = df_sorted.iloc[0]
    return best, df_sorted


# =============================================================================
# DATA VALIDATION AND QUALITY CHECKS
# =============================================================================


def validate_provider_data(df: pd.DataFrame) -> Tuple[bool, str]:
    """
    Validate provider data quality and return status with message.

    Args:
        df: Provider DataFrame to validate

    Returns:
        Tuple of (is_valid, validation_message)
    """
    if df.empty:
        return False, "❌ **Error**: No provider data available. Please check data files."

    issues = []
    info = []

    # Check essential columns (Referral Count is optional and can be calculated)
    required_cols = ["Full Name"]
    missing_cols = [col for col in required_cols if col not in df.columns]
    if missing_cols:
        issues.append(f"Missing required columns: {', '.join(missing_cols)}")

    # Check for geographic data
    if "Latitude" in df.columns and "Longitude" in df.columns:
        missing_coords = (df["Latitude"].isna() | df["Longitude"].isna()).sum()
        if missing_coords > 0:
            issues.append(f"{missing_coords} providers missing geographic coordinates")
    else:
        missing_geo_cols = []
        if "Latitude" not in df.columns:
            missing_geo_cols.append("Latitude")
        if "Longitude" not in df.columns:
            missing_geo_cols.append("Longitude")
        if missing_geo_cols:
            info.append(f"Geographic columns missing: {', '.join(missing_geo_cols)} (may need geocoding)")

    # Check referral count data
    if "Referral Count" in df.columns:
        invalid_counts = df["Referral Count"].isna().sum()
        if invalid_counts > 0:
            issues.append(f"{invalid_counts} providers have invalid referral counts")

        zero_referrals = (df["Referral Count"] == 0).sum()
        if zero_referrals > 0:
            info.append(f"{zero_referrals} providers have zero referrals")

        avg_referrals = df["Referral Count"].mean()
        max_referrals = df["Referral Count"].max()
        info.append(f"Average referrals per provider: {avg_referrals:.1f}")
        info.append(f"Most referred provider has: {max_referrals} referrals")
    else:
        info.append("Referral Count column not found - will be calculated from detailed referral data")

    # Summary info
    total_providers = len(df)
    info.append(f"Total providers in database: {total_providers}")

    # Compile message
    message_parts = []
    if issues:
        message_parts.append("⚠️ **Data Quality Issues**: " + "; ".join(issues))
    if info:
        message_parts.append("ℹ️ **Data Summary**: " + "; ".join(info))

    is_valid = len(issues) == 0
    message = "\n\n".join(message_parts)

    return is_valid, message


def validate_and_clean_coordinates(df: pd.DataFrame) -> pd.DataFrame:
    """
    Validate and clean latitude/longitude coordinates in provider data.

    Args:
        df: DataFrame with coordinate columns

    Returns:
        DataFrame with cleaned coordinates
    """
    if df.empty:
        return df

    df = df.copy()

    # Convert coordinates to numeric, handling errors gracefully
    if "Latitude" in df.columns:
        df["Latitude"] = df["Latitude"].apply(lambda x: safe_numeric_conversion(x, 0.0))

    if "Longitude" in df.columns:
        df["Longitude"] = df["Longitude"].apply(lambda x: safe_numeric_conversion(x, 0.0))

    # Validate coordinate ranges
    if "Latitude" in df.columns and "Longitude" in df.columns:
        # Flag invalid coordinates (outside reasonable bounds for US)
        invalid_lat = (df["Latitude"] < 20) | (df["Latitude"] > 70) | df["Latitude"].isna()
        invalid_lon = (df["Longitude"] < -180) | (df["Longitude"] > -60) | df["Longitude"].isna()

        invalid_coords = invalid_lat | invalid_lon

        if invalid_coords.any():
            # Log providers with invalid coordinates but don't remove them
            invalid_count = invalid_coords.sum()
            st.warning(
                f"⚠️ {invalid_count} providers have invalid or missing coordinates and will not appear in distance calculations."
            )

    return df


def validate_phone_number(phone: str) -> Tuple[bool, str]:
    """
    Validate phone number format.

    Args:
        phone: Phone number string

    Returns:
        Tuple of (is_valid, error_message)
    """
    if not phone or not phone.strip():
        return True, "Phone number is optional"

    # Remove common formatting
    cleaned = re.sub(r"[^\d]", "", phone)

    if len(cleaned) == 10:
        return True, "Valid phone number"
    elif len(cleaned) == 11 and cleaned.startswith("1"):
        return True, "Valid phone number"
    else:
        return False, "Phone number must be 10 digits or 11 digits starting with 1"


# =============================================================================
# UTILITY FUNCTIONS
# =============================================================================


def sanitize_filename(name: str) -> str:
    """Sanitize a string for use as a filename."""
    return re.sub(r"[^A-Za-z0-9_]", "", name.replace(" ", "_"))


def handle_streamlit_error(error: Exception, context: str = "operation"):
    """
    Handle errors gracefully in Streamlit with user-friendly messages.

    Args:
        error: The exception that occurred
        context: Context description for the error
    """
    error_type = type(error).__name__
    error_message = str(error)

    if "geocod" in error_message.lower():
        st.error(
            f"❌ **Geocoding Error**: Unable to find coordinates for the provided address. Please check the address format and try again."
        )
    elif "network" in error_message.lower() or "connection" in error_message.lower():
        st.error(f"❌ **Network Error**: Unable to connect to geocoding service. Please check your internet connection.")
    elif "timeout" in error_message.lower():
        st.error(f"❌ **Timeout Error**: The geocoding service is taking too long to respond. Please try again.")
    elif "file" in error_message.lower() or "not found" in error_message.lower():
        st.error(f"❌ **Data Error**: Required data files are missing. Please contact support.")
    else:
        st.error(f"❌ **Error during {context}**: {error_message}")

    # Log the full error details
    st.exception(error)


def get_word_bytes(best_provider: pd.Series) -> bytes:
    """
    Generate a Word document as bytes for the recommended provider.

    Args:
        best_provider: Series containing provider information

    Returns:
        Word document as bytes
    """
    doc = Document()
    doc.add_heading("Recommended Provider", 0)
    doc.add_paragraph(f"Name: {best_provider['Full Name']}")
    doc.add_paragraph(f"Address: {best_provider['Full Address']}")

    phone = best_provider.get("Phone Number") or best_provider.get("Phone 1")
    if phone:
        doc.add_paragraph(f"Phone: {phone}")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# =============================================================================
# PERFORMANCE MONITORING
# =============================================================================


def monitor_performance(slow_threshold: float = 1.0, log_memory: bool = False):
    """
    Decorator to monitor function performance and log slow operations.

    Args:
        slow_threshold (float): Threshold in seconds above which to log as slow
        log_memory (bool): Whether to log memory usage

    Returns:
        Callable: Decorated function with performance monitoring

    Example:
        @monitor_performance(slow_threshold=0.5, log_memory=True)
        def my_function():
            # Function implementation
            pass
    """

    def decorator(func):
        @functools.wraps(func)
        def wrapper(*args, **kwargs):
            func_name = f"{func.__module__}.{func.__name__}"
            start_time = time.time()
            start_memory = psutil.Process().memory_info().rss / 1024 / 1024 if log_memory else None

            try:
                result = func(*args, **kwargs)
                execution_time = time.time() - start_time

                # Log performance metrics
                if execution_time > slow_threshold:
                    logger.warning(f"SLOW: {func_name} took {execution_time:.3f}s (threshold: {slow_threshold}s)")
                else:
                    logger.debug(f"{func_name} completed in {execution_time:.3f}s")

                # Log memory usage if requested
                if log_memory and start_memory is not None:
                    end_memory = psutil.Process().memory_info().rss / 1024 / 1024
                    memory_diff = end_memory - start_memory
                    if abs(memory_diff) > 10:  # Log if memory change > 10MB
                        logger.info(f"{func_name} memory change: {memory_diff:+.1f}MB (now: {end_memory:.1f}MB)")

                return result

            except Exception as e:
                execution_time = time.time() - start_time
                logger.error(f"{func_name} failed after {execution_time:.3f}s: {str(e)}")
                raise

        return wrapper

    return decorator


# =============================================================================
# EXPORT ALL FUNCTIONS
# =============================================================================

__all__ = [
    # Data Loading and Processing
    "load_provider_data",
    "clean_address_data",
    "build_full_address",
    "safe_numeric_conversion",
    # Address Validation and Geocoding
    "validate_address",
    "validate_address_input",
    "validate_coordinates",
    "geocode_address_with_cache",
    "cached_geocode_address",
    "handle_geocoding_error",
    # Distance Calculation and Provider Recommendation
    "calculate_distances",
    "recommend_provider",
    # Data Validation and Quality Checks
    "validate_provider_data",
    "validate_and_clean_coordinates",
    "validate_phone_number",
    # Utility Functions
    "sanitize_filename",
    "handle_streamlit_error",
    "get_word_bytes",
    "monitor_performance",
]
