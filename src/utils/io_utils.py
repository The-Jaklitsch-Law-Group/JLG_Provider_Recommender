"""IO and small helpers: docx export, filename sanitization, and streamlit error handler."""
import io
import re

import pandas as pd
import streamlit as st
from docx import Document


def format_phone_number(phone):
    """
    Convert phone number to formatted string "(XXX) XXX-XXXX".
    Handles various input types including float, int, and string.

    Args:
        phone: Phone number as float, int, or string

    Returns:
        Formatted phone string or original value if formatting fails
    """
    if pd.isna(phone) or phone is None:
        return None

    # Handle different input types
    if isinstance(phone, float):
        phone = int(phone)
    elif isinstance(phone, str) and '.' in phone:
        # Handle strings that look like floats (e.g., '4435140560.0')
        try:
            phone = int(float(phone))
        except (ValueError, TypeError):
            pass  # Keep as string if conversion fails

    # Convert to string and remove any non-digit characters
    phone_str = str(phone).replace('-', '').replace('(', '').replace(')', '').replace(' ', '')

    # Extract only digits
    digits = ''.join(filter(str.isdigit, phone_str))

    # Check if we have exactly 10 digits
    if len(digits) == 10:
        return f"({digits[:3]}) {digits[3:6]}-{digits[6:]}"
    elif len(digits) == 11 and digits[0] == '1':
        # Handle numbers with leading 1
        return f"({digits[1:4]}) {digits[4:7]}-{digits[7:]}"
    else:
        # Return original if we can't format it properly
        return phone


def get_word_bytes(best_provider: pd.Series) -> bytes:
    doc = Document()
    doc.add_heading("Recommended Provider", 0)
    doc.add_paragraph(f"Name: {best_provider.get('Full Name', '')}")
    # Preferred Provider status if available (display as Yes/No for booleans)
    pref = best_provider.get("Preferred Provider")
    if pref is not None:
        if isinstance(pref, bool):
            pref_display = "Yes" if pref else "No"
        else:
            pref_display = str(pref)
        doc.add_paragraph(f"Preferred Provider: {pref_display}")
    doc.add_paragraph(f"Address: {best_provider.get('Full Address', '')}")
    phone = None
    for phone_key in ["Work Phone Number", "Work Phone", "Phone Number", "Phone 1"]:
        candidate = best_provider.get(phone_key)
        if candidate:
            phone = format_phone_number(candidate)
            break
    if phone:
        doc.add_paragraph(f"Phone: {phone}")
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def sanitize_filename(name: str) -> str:
    return re.sub(r"[^A-Za-z0-9_]", "", name.replace(" ", "_"))


def handle_streamlit_error(error: Exception, context: str = "operation") -> None:
    err = str(error)
    if "geocod" in err.lower():
        st.error(
            (
                "❌ **Geocoding Error**: Unable to find coordinates for the provided address. "
                "Please check the address format and try again."
            )
        )
    elif "network" in err.lower() or "connection" in err.lower():
        st.error("❌ **Network Error**: Unable to connect to geocoding service. Please check your internet connection.")
    elif "timeout" in err.lower():
        st.error("❌ **Timeout Error**: The geocoding service is taking too long to respond. Please try again.")
    elif "file" in err.lower() or "not found" in err.lower():
        st.error("❌ **Data Error**: Required data files are missing. Please contact support.")
    else:
        st.error(f"❌ **Error during {context}**: {err}")

    st.exception(error)
