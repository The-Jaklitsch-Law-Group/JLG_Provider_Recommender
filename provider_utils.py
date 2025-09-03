import io
import logging
import re
import time
from datetime import datetime, timedelta
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import numpy as np
import pandas as pd
import streamlit as st
from docx import Document
from geopy.distance import geodesic
from geopy.exc import GeocoderServiceError, GeocoderTimedOut, GeocoderUnavailable
from geopy.geocoders import Nominatim


# --- Data Loading ---
@st.cache_data(ttl=3600)
def load_provider_data(filepath: str) -> pd.DataFrame:
    """
    Load and preprocess provider data from various file formats.

    This function loads healthcare provider data from Excel, CSV, Feather, or Parquet files
    and performs basic preprocessing including column cleaning, data type conversion,
    and address formatting.

    Args:
        filepath (str): Path to the data file. Supported formats: .xlsx, .csv, .feather, .parquet

    Returns:
        pd.DataFrame: Preprocessed provider data with columns:
            - Full Name: Provider's complete name
            - Specialty: Medical specialty
            - Street, City, State, Zip: Address components
            - Full Address: Formatted complete address
            - Referral Count: Number of referrals (converted to numeric)
            - Latitude, Longitude: Geographic coordinates (if available)

    Raises:
        FileNotFoundError: If the specified file does not exist
        ValueError: If the file format is not supported

    Example:
        >>> df = load_provider_data("data/providers.xlsx")
        >>> print(df.columns.tolist())
        ['Full Name', 'Specialty', 'Street', 'City', 'State', 'Zip', 'Full Address', 'Referral Count']

    Note:
        - Function is cached with Streamlit for 1 hour to improve performance
        - Preference column is automatically dropped if present
        - ZIP codes are converted to strings to preserve leading zeros
        - Missing address components are filled with empty strings
    """

    path = Path(filepath)
    if not path.exists():
        raise FileNotFoundError(f"File {filepath} does not exist")

    suffix = path.suffix.lower()
    if suffix == ".xlsx":
        df = pd.read_excel(path)
    elif suffix == ".csv":
        df = pd.read_csv(path)
    elif suffix == ".feather":
        df = pd.read_feather(path)
    elif suffix == ".parquet":
        df = pd.read_parquet(path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")

    df.columns = [col.strip() for col in df.columns]
    df = df.drop(columns="Preference", errors="ignore")
    df["Zip"] = df["Zip"].apply(lambda x: str(x) if pd.notnull(x) else "")
    df["Referral Count"] = pd.to_numeric(df["Referral Count"], errors="coerce")
    df["Full Address"] = (
        df["Street"].fillna("")
        + ", "
        + df["City"].fillna("")
        + ", "
        + df["State"].fillna("")
        + " "
        + df["Zip"].fillna("")
    )
    df["Full Address"] = df["Full Address"].str.replace(r",\s*,", ",", regex=True).str.replace(r",\s*$", "", regex=True)
    return df


@st.cache_data(ttl=3600)
def load_detailed_referrals(filepath: str) -> pd.DataFrame:
    """Load detailed referral data with dates for time-based filtering."""

    path = Path(filepath)
    if not path.exists():
        # If detailed data doesn't exist, return empty DataFrame
        return pd.DataFrame()

    try:
        df = pd.read_parquet(path)
        df["Referral Date"] = pd.to_datetime(df["Referral Date"], errors="coerce")
        return df
    except Exception as e:
        st.warning(f"Could not load detailed referral data: {e}")
        return pd.DataFrame()


def calculate_time_based_referral_counts(detailed_df: pd.DataFrame, start_date, end_date) -> pd.DataFrame:
    """Calculate referral counts for providers within a specific time period."""

    if detailed_df.empty:
        return pd.DataFrame()

    # Filter by date range
    if start_date and end_date:
        mask = (detailed_df["Referral Date"] >= pd.to_datetime(start_date)) & (
            detailed_df["Referral Date"] <= pd.to_datetime(end_date)
        )
        filtered_df = detailed_df[mask]
    else:
        filtered_df = detailed_df

    if filtered_df.empty:
        return pd.DataFrame()

    # Group by provider and count referrals
    provider_cols = [
        "Person ID",
        "Full Name",
        "Street",
        "City",
        "State",
        "Zip",
        "Latitude",
        "Longitude",
        "Phone Number",
    ]

    # Only include columns that exist in the DataFrame
    available_cols = [col for col in provider_cols if col in filtered_df.columns]

    if not available_cols:
        return pd.DataFrame()

    time_based_counts = (
        filtered_df.groupby(available_cols, as_index=False)
        .size()
        .rename(columns={"size": "Referral Count"})
        .sort_values(by="Referral Count", ascending=False)
    )

    # Add Full Address if not present
    if "Full Address" not in time_based_counts.columns and all(
        col in time_based_counts.columns for col in ["Street", "City", "State", "Zip"]
    ):
        time_based_counts["Full Address"] = (
            time_based_counts["Street"].fillna("")
            + ", "
            + time_based_counts["City"].fillna("")
            + ", "
            + time_based_counts["State"].fillna("")
            + " "
            + time_based_counts["Zip"].fillna("")
        )
        time_based_counts["Full Address"] = (
            time_based_counts["Full Address"]
            .str.replace(r",\s*,", ",", regex=True)
            .str.replace(r",\s*$", "", regex=True)
        )

    return time_based_counts


def safe_numeric_conversion(value: Any, default: float = 0.0) -> float:
    """
    Safely convert a value to float with fallback to default.

    This function handles common conversion errors when dealing with mixed data types,
    missing values, or malformed numeric strings. It's particularly useful for
    cleaning data from external sources.

    Args:
        value (Any): Value to convert to float. Can be str, int, float, None, or pandas.NA
        default (float, optional): Default value to return if conversion fails. Defaults to 0.0

    Returns:
        float: Successfully converted float value or the default value

    Examples:
        >>> safe_numeric_conversion("123.45")
        123.45
        >>> safe_numeric_conversion("invalid", default=0.0)
        0.0
        >>> safe_numeric_conversion(None)
        0.0
        >>> safe_numeric_conversion(pd.NA, default=-1.0)
        -1.0

    Note:
        Uses pandas.isna() to check for missing values, which handles both
        None and pandas NA values properly.
    """
    try:
        if pd.isna(value):
            return default
        return float(value)
    except (ValueError, TypeError):
        return default


def load_and_validate_provider_data(
    start_date: Optional[datetime] = None, end_date: Optional[datetime] = None
) -> pd.DataFrame:
    """
    Load provider data with optional time filtering and validation.

    Args:
        start_date: Start date for filtering referrals
        end_date: End date for filtering referrals

    Returns:
        Validated provider dataframe
    """
    try:
        # Try to load the detailed referrals data first
        try:
            df = pd.read_parquet("data/detailed_referrals.parquet")

            # Convert date columns if they exist
            if "Referral Date" in df.columns:
                df["Referral Date"] = pd.to_datetime(df["Referral Date"], errors="coerce")

                # Apply time filtering if dates provided
                if start_date and end_date:
                    mask = (df["Referral Date"] >= start_date) & (df["Referral Date"] <= end_date)
                    df = df[mask]

                    # Group by provider and count referrals - adapt to actual column names
                    groupby_cols = []
                    for col in ["Full Name", "Street", "City", "State", "Zip", "Latitude", "Longitude", "Phone Number"]:
                        if col in df.columns:
                            groupby_cols.append(col)

                    if groupby_cols:
                        referral_counts = df.groupby(groupby_cols).size().reset_index(name="Referral Count")
                        df = referral_counts
                else:
                    # Use pre-aggregated data if no time filtering
                    if "Referral Count" not in df.columns:
                        groupby_cols = []
                        for col in [
                            "Full Name",
                            "Street",
                            "City",
                            "State",
                            "Zip",
                            "Latitude",
                            "Longitude",
                            "Phone Number",
                        ]:
                            if col in df.columns:
                                groupby_cols.append(col)
                        if groupby_cols:
                            df = df.groupby(groupby_cols).size().reset_index(name="Referral Count")

        except FileNotFoundError:
            # Fallback to original cleaned data
            st.warning("Detailed referrals data not found. Using aggregated data (time filtering not available).")
            df = pd.read_parquet("data/cleaned_outbound_referrals.parquet")

        # Validate data
        is_valid, issues = validate_provider_data(df)
        if not is_valid and isinstance(issues, list):
            st.warning(f"Data quality issues detected: {'; '.join(issues)}")

        # Clean and standardize coordinates
        if "Latitude" in df.columns:
            df["Latitude"] = df["Latitude"].apply(lambda x: safe_numeric_conversion(x, 0.0))
        if "Longitude" in df.columns:
            df["Longitude"] = df["Longitude"].apply(lambda x: safe_numeric_conversion(x, 0.0))

        # Remove providers with invalid coordinates
        if "Latitude" in df.columns and "Longitude" in df.columns:
            df = df[(df["Latitude"] != 0) & (df["Longitude"] != 0)]

        return df

    except Exception as e:
        st.error(f"Error loading provider data: {str(e)}")
        return pd.DataFrame()


def handle_streamlit_error(error: Exception, context: str = "operation"):
    """
    Handle errors gracefully in Streamlit with user-friendly messages.

    Args:
        error: The exception that occurred
        context: Context description for the error
    """
    error_type = type(error).__name__
    error_message = str(error)

    if "geocod" in error_message.lower():
        st.error(
            f"❌ **Geocoding Error**: Unable to find coordinates for the provided address. Please check the address format and try again."
        )
    elif "network" in error_message.lower() or "connection" in error_message.lower():
        st.error(f"❌ **Network Error**: Unable to connect to geocoding service. Please check your internet connection.")
    elif "timeout" in error_message.lower():
        st.error(f"❌ **Timeout Error**: The geocoding service is taking too long to respond. Please try again.")
    elif "file" in error_message.lower() or "not found" in error_message.lower():
        st.error(f"❌ **Data Error**: Required data files are missing. Please contact support.")
    else:
        st.error(f"❌ **Error during {context}**: {error_message}")

    # Log the full error details
    st.exception(error)


@st.cache_data(ttl=3600)
def geocode_address_with_cache(address: str) -> Optional[Tuple[float, float]]:
    """
    Geocode an address with caching and error handling.

    Args:
        address: Address string to geocode

    Returns:
        Tuple of (latitude, longitude) or None if failed
    """
    try:
        geolocator = Nominatim(user_agent="jlg_provider_recommender", timeout=10)
        location = geolocator.geocode(address)

        if location:
            return (location.latitude, location.longitude)
        else:
            return None

    except (GeocoderTimedOut, GeocoderServiceError) as e:
        st.warning(f"Geocoding service temporarily unavailable. Please try again.")
        return None
    except Exception as e:
        st.error(f"Error geocoding address: {str(e)}")
        return None


def validate_address(address: str) -> Tuple[bool, str]:
    """
    Validate address format and completeness.

    Args:
        address: Address string to validate

    Returns:
        Tuple of (is_valid, validation_message)
    """
    if not address or not address.strip():
        return False, "Address cannot be empty"

    address = address.strip()

    # Basic length check
    if len(address) < 10:
        return False, "Address appears too short. Please provide a complete address."

    # Check for basic components
    has_number = any(char.isdigit() for char in address)
    has_street_indicator = any(
        indicator in address.lower()
        for indicator in [
            "street",
            "st",
            "avenue",
            "ave",
            "road",
            "rd",
            "drive",
            "dr",
            "lane",
            "ln",
            "way",
            "blvd",
            "boulevard",
        ]
    )

    if not has_number:
        return False, "Address should include a street number"

    # Check for state abbreviation or full state name
    state_patterns = ["\\b[A-Z]{2}\\b", "\\d{5}(-\\d{4})?\\b"]  # State abbrev or ZIP
    has_state_or_zip = any(re.search(pattern, address) for pattern in state_patterns)

    if not has_state_or_zip:
        return True, "Consider adding state and ZIP code for better accuracy"

    return True, ""


def sanitize_filename(name):
    """Sanitize a string for use as a filename."""
    return re.sub(r"[^A-Za-z0-9_]", "", name.replace(" ", "_"))


def validate_address_input(street: str, city: str, state: str, zipcode: str) -> tuple[bool, str]:
    """
    Comprehensive address validation with enhanced security and validation.

    This function validates individual address components using enhanced validation
    rules and security checks to ensure data integrity and prevent injection attacks.

    Args:
        street (str): Street address (e.g., "123 Main St")
        city (str): City name (e.g., "New York")
        state (str): State abbreviation (e.g., "NY")
        zipcode (str): ZIP code (e.g., "10001" or "10001-1234")

    Returns:
        tuple[bool, str]: (is_valid, detailed_error_message)

    Examples:
        >>> is_valid, msg = validate_address_input("123 Main St", "Anytown", "CA", "12345")
        >>> print(is_valid)  # True

        >>> is_valid, msg = validate_address_input("", "Anytown", "INVALID", "12345")
        >>> print(msg)  # "Street address must be at least 3 characters; State must be a valid 2-letter abbreviation"
    """
    try:
        # Import validation utilities
        from security_utils import InputValidator

        # Use enhanced validation
        return InputValidator.validate_address_input(street, city, state, zipcode)

    except ImportError:
        # Fallback to basic validation if security_utils not available
        pass

    errors = []
    warnings = []

    # Basic required field validation
    if not street.strip():
        errors.append("Street address is required")

    if not city.strip():
        warnings.append("City is recommended for better geocoding accuracy")

    if not state.strip():
        warnings.append("State is recommended for better geocoding accuracy")

    # Validate state format (if provided)
    if state.strip():
        state_clean = state.strip().upper()
        # Common US state abbreviations
        valid_states = {
            "AL",
            "AK",
            "AZ",
            "AR",
            "CA",
            "CO",
            "CT",
            "DE",
            "FL",
            "GA",
            "HI",
            "ID",
            "IL",
            "IN",
            "IA",
            "KS",
            "KY",
            "LA",
            "ME",
            "MD",
            "MA",
            "MI",
            "MN",
            "MS",
            "MO",
            "MT",
            "NE",
            "NV",
            "NH",
            "NJ",
            "NM",
            "NY",
            "NC",
            "ND",
            "OH",
            "OK",
            "OR",
            "PA",
            "RI",
            "SC",
            "SD",
            "TN",
            "TX",
            "UT",
            "VT",
            "VA",
            "WA",
            "WV",
            "WI",
            "WY",
            "DC",
        }
        if len(state_clean) == 2 and state_clean not in valid_states:
            warnings.append(f"'{state}' may not be a valid US state abbreviation")
        elif len(state_clean) > 2:
            warnings.append("Consider using 2-letter state abbreviation (e.g., 'MD' instead of 'Maryland')")

    # Validate ZIP code format (if provided)
    if zipcode.strip():
        zip_clean = zipcode.strip()
        if not (zip_clean.isdigit() and len(zip_clean) == 5):
            if not (
                len(zip_clean) == 10 and zip_clean[5] == "-" and zip_clean[:5].isdigit() and zip_clean[6:].isdigit()
            ):
                warnings.append("ZIP code should be 5 digits (e.g., '20746') or ZIP+4 format (e.g., '20746-1234')")

    # Check for suspicious patterns
    if street.strip().lower() in ["test", "example", "123 test st", "123 main st"]:
        warnings.append("Address appears to be a test value - please enter a real address")

    # Compile messages
    message_parts = []
    if errors:
        message_parts.append("❌ **Errors**: " + "; ".join(errors))
    if warnings:
        message_parts.append("⚠️ **Suggestions**: " + "; ".join(warnings))

    is_valid = len(errors) == 0
    message = "\n\n".join(message_parts) if message_parts else ""

    return is_valid, message


def handle_geocoding_error(address: str, error: Exception) -> str:
    """Provide user-friendly error messages for geocoding failures."""

    error_type = type(error).__name__

    if "timeout" in str(error).lower():
        return f"⏱️ **Geocoding Timeout**: The address lookup service is taking too long. Please try again in a moment."

    elif "unavailable" in str(error).lower() or "service" in str(error).lower():
        return f"🔌 **Service Unavailable**: The geocoding service is temporarily unavailable. Please try again later."

    elif "rate" in str(error).lower() or "limit" in str(error).lower():
        return f"🚦 **Rate Limited**: Too many requests to the geocoding service. Please wait a moment and try again."

    elif "network" in str(error).lower() or "connection" in str(error).lower():
        return f"🌐 **Network Error**: Cannot connect to the geocoding service. Please check your internet connection."

    else:
        return f"❌ **Geocoding Error**: Unable to find location for '{address}'. Please check the address and try again. (Error: {error_type})"


def validate_and_clean_coordinates(df: pd.DataFrame) -> pd.DataFrame:
    """Validate and clean latitude/longitude coordinates in provider data."""

    if df.empty:
        return df

    df = df.copy()

    # Convert coordinates to numeric, handling errors gracefully
    if "Latitude" in df.columns:
        df["Latitude"] = df["Latitude"].apply(lambda x: safe_numeric_conversion(x, 0.0))

    if "Longitude" in df.columns:
        df["Longitude"] = df["Longitude"].apply(lambda x: safe_numeric_conversion(x, 0.0))

    # Validate coordinate ranges
    if "Latitude" in df.columns and "Longitude" in df.columns:
        # Flag invalid coordinates (outside reasonable bounds for US)
        invalid_lat = (df["Latitude"] < 20) | (df["Latitude"] > 70) | df["Latitude"].isna()
        invalid_lon = (df["Longitude"] < -180) | (df["Longitude"] > -60) | df["Longitude"].isna()

        invalid_coords = invalid_lat | invalid_lon

        if invalid_coords.any():
            # Log providers with invalid coordinates but don't remove them
            invalid_count = invalid_coords.sum()
            st.warning(
                f"⚠️ {invalid_count} providers have invalid or missing coordinates and will not appear in distance calculations."
            )

    return df


def validate_provider_data(df: pd.DataFrame) -> tuple[bool, str]:
    """Validate provider data quality and return status with message."""

    if df.empty:
        return False, "❌ **Error**: No provider data available. Please check data files."

    issues = []
    info = []

    # Check required columns
    required_cols = ["Full Name", "Referral Count", "Latitude", "Longitude"]
    missing_cols = [col for col in required_cols if col not in df.columns]
    if missing_cols:
        issues.append(f"Missing required columns: {', '.join(missing_cols)}")

    # Check data quality
    if "Referral Count" in df.columns:
        invalid_counts = df["Referral Count"].isna().sum()
        if invalid_counts > 0:
            issues.append(f"{invalid_counts} providers have invalid referral counts")

        zero_referrals = (df["Referral Count"] == 0).sum()
        if zero_referrals > 0:
            info.append(f"{zero_referrals} providers have zero referrals")

    if "Latitude" in df.columns and "Longitude" in df.columns:
        missing_coords = (df["Latitude"].isna() | df["Longitude"].isna()).sum()
        if missing_coords > 0:
            issues.append(f"{missing_coords} providers missing geographic coordinates")

    # Summary info
    total_providers = len(df)
    info.append(f"Total providers in database: {total_providers}")

    if "Referral Count" in df.columns:
        avg_referrals = df["Referral Count"].mean()
        max_referrals = df["Referral Count"].max()
        info.append(f"Average referrals per provider: {avg_referrals:.1f}")
        info.append(f"Most referred provider has: {max_referrals} referrals")

    # Compile message
    message_parts = []
    if issues:
        message_parts.append("⚠️ **Data Quality Issues**: " + "; ".join(issues))
    if info:
        message_parts.append("ℹ️ **Data Summary**: " + "; ".join(info))

    is_valid = len(issues) == 0
    message = "\n\n".join(message_parts)

    return is_valid, message
    """Validate provider data quality and return status with message."""

    if df.empty:
        return False, "❌ **Error**: No provider data available. Please check data files."

    issues = []
    info = []

    # Check required columns
    required_cols = ["Full Name", "Referral Count", "Latitude", "Longitude"]
    missing_cols = [col for col in required_cols if col not in df.columns]
    if missing_cols:
        issues.append(f"Missing required columns: {', '.join(missing_cols)}")

    # Check data quality
    if "Referral Count" in df.columns:
        invalid_counts = df["Referral Count"].isna().sum()
        if invalid_counts > 0:
            issues.append(f"{invalid_counts} providers have invalid referral counts")

        zero_referrals = (df["Referral Count"] == 0).sum()
        if zero_referrals > 0:
            info.append(f"{zero_referrals} providers have zero referrals")

    if "Latitude" in df.columns and "Longitude" in df.columns:
        missing_coords = (df["Latitude"].isna() | df["Longitude"].isna()).sum()
        if missing_coords > 0:
            issues.append(f"{missing_coords} providers missing geographic coordinates")

    # Summary info
    total_providers = len(df)
    info.append(f"Total providers in database: {total_providers}")

    if "Referral Count" in df.columns:
        avg_referrals = df["Referral Count"].mean()
        max_referrals = df["Referral Count"].max()
        info.append(f"Average referrals per provider: {avg_referrals:.1f}")
        info.append(f"Most referred provider has: {max_referrals} referrals")

    # Compile message
    message_parts = []
    if issues:
        message_parts.append("⚠️ **Data Quality Issues**: " + "; ".join(issues))
    if info:
        message_parts.append("ℹ️ **Data Summary**: " + "; ".join(info))

    is_valid = len(issues) == 0
    message = "\n\n".join(message_parts)

    return is_valid, message


def get_word_bytes(best):
    """Generate a Word document as bytes for the recommended provider."""
    doc = Document()
    doc.add_heading("Recommended Provider", 0)
    doc.add_paragraph(f"Name: {best['Full Name']}")
    doc.add_paragraph(f"Address: {best['Full Address']}")
    phone = best.get("Phone Number") or best.get("Phone 1")
    if phone:
        doc.add_paragraph(f"Phone: {phone}")
    # doc.add_paragraph(f"Email: {best['Email 1']}")
    # doc.add_paragraph(f"Specialty: {best['Specialty']}")
    # if best.get('Preferred', 0) == 1:
    #     doc.add_paragraph("Preferred Provider")
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    # Return raw bytes which can be passed directly to Streamlit's download APIs
    return buffer.getvalue()


def recommend_provider(provider_df, distance_weight=0.5, referral_weight=0.5, min_referrals=None):
    """
    Recommend the best healthcare provider based on distance and referral score.

    This function implements a weighted scoring algorithm that combines provider distance
    from the client and historical referral patterns to recommend the most suitable
    healthcare provider. The algorithm normalizes both metrics and allows customizable
    weighting between proximity and provider popularity.

    Args:
        provider_df (pd.DataFrame): DataFrame containing provider information with required columns:
            - 'Distance (Miles)': Float distance from client to provider
            - 'Referral Count': Integer number of historical referrals
            - 'Full Name': Provider name
            - Additional columns preserved in output
        distance_weight (float, optional): Weight for distance factor (0.0-1.0). Defaults to 0.5.
            Higher values prioritize closer providers.
        referral_weight (float, optional): Weight for referral count factor (0.0-1.0). Defaults to 0.5.
            Higher values prioritize more frequently referred providers.
        min_referrals (int, optional): Minimum referral count threshold. Providers with fewer
            referrals are excluded. Defaults to None (no filtering).

    Returns:
        tuple: A tuple containing:
            - best_provider (pd.Series or None): Series containing the top recommended provider's data
            - scored_df (pd.DataFrame or None): DataFrame with all eligible providers and their scores,
              sorted by recommendation score (lower is better)

    Raises:
        KeyError: If required columns ('Distance (Miles)', 'Referral Count') are missing

    Algorithm:
        1. Filters out providers with missing distance or referral data
        2. Applies minimum referral threshold if specified
        3. Normalizes distance (0-1, lower is better) and referral count (0-1, higher is better)
        4. Calculates composite score: distance_weight * norm_distance + referral_weight * (1 - norm_referrals)
        5. Returns provider with lowest composite score

    Examples:
        >>> # Basic usage with equal weighting
        >>> best, scored = recommend_provider(providers_df)

        >>> # Prioritize distance over referral history
        >>> best, scored = recommend_provider(providers_df, distance_weight=0.8, referral_weight=0.2)

        >>> # Only consider providers with at least 5 referrals
        >>> best, scored = recommend_provider(providers_df, min_referrals=5)

    Note:
        - Returns (None, None) if no providers meet the criteria
        - Weights don't need to sum to 1.0, but typical usage has distance_weight + referral_weight = 1.0
        - Lower composite scores indicate better recommendations
        - Preferred providers can be prioritized by uncommenting the preferred provider logic
    """
    df = provider_df.copy(deep=True)
    df = df[df["Distance (Miles)"].notnull() & df["Referral Count"].notnull()]
    if min_referrals is not None:
        df = df[df["Referral Count"] >= min_referrals]

    if df.empty:
        return None, None

    # # Prioritize preferred providers: filter to preferred if any exist
    # preferred_df = df[df['Preferred'] == 1]
    # if not preferred_df.empty:
    #     df = preferred_df

    # Safe normalization (avoid division by zero)
    referral_range = df["Referral Count"].max() - df["Referral Count"].min()
    dist_range = df["Distance (Miles)"].max() - df["Distance (Miles)"].min()
    df["norm_rank"] = (df["Referral Count"] - df["Referral Count"].min()) / referral_range if referral_range != 0 else 0
    df["norm_dist"] = (df["Distance (Miles)"] - df["Distance (Miles)"].min()) / dist_range if dist_range != 0 else 0
    df["Score"] = distance_weight * df["norm_dist"] + referral_weight * df["norm_rank"]

    # Tie-break behavior:
    # - If distance is prioritized (distance_weight > referral_weight), sort ties by
    #   Distance (ascending) then Referral Count (ascending).
    # - Otherwise, sort ties by Referral Count (ascending) then Distance (ascending).
    # This ensures deterministic ordering when multiple providers share the same score.
    if distance_weight > referral_weight:
        sort_keys = ["Score", "Distance (Miles)", "Referral Count"]
    else:
        sort_keys = ["Score", "Referral Count", "Distance (Miles)"]

    # Add a stable final tie-break key so ordering is deterministic across runs/users.
    # Prefer alphabetical provider name if available.
    candidate_keys = sort_keys + ["Full Name"]
    sort_keys_final = [k for k in candidate_keys if k in df.columns]
    ascending = [True] * len(sort_keys_final)

    df_sorted = df.sort_values(by=sort_keys_final, ascending=ascending).reset_index(drop=True)
    best = df_sorted.iloc[0]
    return best, df_sorted


_geocode_cache = {}


def geocode_address(addresses, _geocode):
    """Geocode a list of addresses, returning lists of latitudes and longitudes."""

    def _cached_geocode(addr):
        if addr not in _geocode_cache:
            try:
                location = _geocode(addr, timeout=10)
                if location:
                    _geocode_cache[addr] = (location.latitude, location.longitude)
                else:
                    _geocode_cache[addr] = (None, None)
            except GeocoderUnavailable as e:
                _geocode_cache[addr] = (None, None)
                print(f"GeocoderUnavailable for address '{addr}': {e}")
            except Exception as e:
                _geocode_cache[addr] = (None, None)
                print(f"Geocoding error for address '{addr}': {e}")
        return _geocode_cache[addr]

    results = pd.Series(addresses).map(_cached_geocode)
    lats, lons = zip(*results)
    return list(lats), list(lons)


@st.cache_data(ttl=60 * 60 * 24)
def cached_geocode_address(q: str):
    """Cached single-address geocode using Nominatim+RateLimiter.

    Returns a geopy Location or None. TTL 24 hours to reuse results.
    """
    try:
        from geopy.extra.rate_limiter import RateLimiter
        from geopy.geocoders import Nominatim

        geolocator_local = Nominatim(user_agent="provider_recommender")
        geocode_local = RateLimiter(geolocator_local.geocode, min_delay_seconds=2, max_retries=3)
        return geocode_local(q, timeout=10)
    except GeocoderUnavailable:
        return None
    except Exception:
        return None


def calculate_distances(user_lat, user_lon, provider_df):
    """Calculate distances in miles from user to each provider using vectorized operations."""

    lat_rad = np.radians(provider_df["Latitude"].to_numpy(dtype=float))
    lon_rad = np.radians(provider_df["Longitude"].to_numpy(dtype=float))
    user_lat_rad = np.radians(user_lat)
    user_lon_rad = np.radians(user_lon)

    valid = ~np.isnan(lat_rad) & ~np.isnan(lon_rad)
    dlat = lat_rad[valid] - user_lat_rad
    dlon = lon_rad[valid] - user_lon_rad
    a = np.sin(dlat / 2) ** 2 + np.cos(user_lat_rad) * np.cos(lat_rad[valid]) * np.sin(dlon / 2) ** 2
    c = 2 * np.arcsin(np.sqrt(a))
    distances = np.full(len(provider_df), np.nan)
    distances[valid] = 3958.8 * c  # Earth radius in miles

    # Convert NaN to None for consistency with previous behavior
    return [None if np.isnan(d) else float(d) for d in distances]
